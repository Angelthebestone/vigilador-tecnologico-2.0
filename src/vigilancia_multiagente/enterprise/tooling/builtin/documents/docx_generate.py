"""docx_generate — Tool Tier 1 para generación de DOCX (FR-020).

Genera documentos DOCX desde datos estructurados usando python-docx.
Implementa el protocolo ToolWrapper.
"""

from __future__ import annotations

from pathlib import Path
from typing import Any

from vigilancia_multiagente.enterprise.tooling.tool_wrapper import HealthcheckResult


class DocxGenerateTool:
    """Genera documentos DOCX desde título + secciones estructuradas."""

    name: str = "docx_generate"
    domain: str = "documents"
    is_external_mcp: bool = False
    requires_auth: bool = False

    async def healthcheck(self) -> HealthcheckResult:
        """Verifica que python-docx está disponible."""
        try:
            import docx  # noqa: F401

            return HealthcheckResult(status="UP")
        except ImportError:
            return HealthcheckResult(status="DOWN", error="python-docx not installed")

    async def execute(self, tool_name: str, args: dict[str, Any]) -> dict[str, object]:
        """Genera un DOCX desde datos estructurados.

        Args esperados en `args`:
            title: str — título del documento.
            sections: list[dict] — cada dict con 'heading' y 'body'.
            template_path: str | None — template DOCX base (opcional).
            output_path: str — ruta destino del archivo generado.

        Returns:
            dict con output_path (str) y page_count (int).
        """
        title = args.get("title")
        if not title or not isinstance(title, str):
            return {"error": "Missing or invalid 'title' argument"}

        sections = args.get("sections")
        if not isinstance(sections, list) or len(sections) == 0:
            return {
                "error": "'sections' must be a non-empty list of dicts with 'heading' and 'body'"
            }

        output_path = args.get("output_path")
        if not output_path or not isinstance(output_path, str):
            return {"error": "Missing or invalid 'output_path' argument"}

        for idx, section in enumerate(sections):
            if not isinstance(section, dict):
                return {"error": f"Section at index {idx} must be a dict"}
            if "heading" not in section or "body" not in section:
                return {"error": f"Section at index {idx} must have 'heading' and 'body' keys"}

        try:
            from docx import Document
        except ImportError:
            return {"error": "python-docx is not installed; cannot generate DOCX"}

        template_path = args.get("template_path")
        if template_path and Path(template_path).exists():
            doc = Document(template_path)
        else:
            doc = Document()

        doc.add_heading(title, level=0)

        for section in sections:
            doc.add_heading(str(section["heading"]), level=1)
            doc.add_paragraph(str(section["body"]))

        dest = Path(output_path)
        dest.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(dest))

        # python-docx no expone page_count directamente; estimamos por secciones
        page_count = max(1, len(sections) // 3)

        return {"output_path": str(dest), "page_count": page_count}
